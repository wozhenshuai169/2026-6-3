from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "参赛提交文档"
ASSET_DIR = OUT / "assets"
CAPTURE_DIR = ROOT / "deliverables" / "competition-video" / "captures-latest"

DATE = "2026年7月19日"
VERSION = "V1.0"
TEAM = "云游智导项目组"
PRODUCT = "云游智导——景区智能导览与三端协同系统"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "223047"
ORANGE = "C65D3B"
TEAL = "34765F"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_ORANGE = "FAEEE8"
LIGHT_GREEN = "EAF4F0"
MUTED = "667085"
INK = "000000"
WHITE = "FFFFFF"
LINE = "D8DEE8"

FONT_CJK_BODY = "宋体"
FONT_CJK_HEADING = "黑体"
FONT_CJK_KAI = "楷体"
FONT_LATIN = "Consolas"

FONT_BODY_FILE = Path(r"C:\Windows\Fonts\simsun.ttc")
FONT_HEADING_FILE = Path(r"C:\Windows\Fonts\simhei.ttf")
FONT_KAI_FILE = Path(r"C:\Windows\Fonts\simkai.ttf")
FONT_LATIN_FILE = Path(r"C:\Windows\Fonts\consola.ttf")
FONT_LATIN_BOLD_FILE = Path(r"C:\Windows\Fonts\consolab.ttf")


class MixedFont:
    def __init__(self, size: int, bold: bool = False, cjk: str = "body") -> None:
        self.size = size
        cjk_path = {
            "body": FONT_BODY_FILE,
            "heading": FONT_HEADING_FILE,
            "kai": FONT_KAI_FILE,
        }[cjk]
        self.cjk = ImageFont.truetype(str(cjk_path), size=size)
        self.latin = ImageFont.truetype(str(FONT_LATIN_BOLD_FILE if bold else FONT_LATIN_FILE), size=size)


def font(size: int, bold: bool = False, cjk: str | None = None) -> MixedFont:
    return MixedFont(size=size, bold=bold, cjk=cjk or ("heading" if bold else "body"))


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_east_asia(run, name: str = FONT_CJK_BODY) -> None:
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT_LATIN)


def shade(element, fill: str) -> None:
    shd = element.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        element.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_field(paragraph, instruction: str, display: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_numbering(doc: Document, kind: str, levels: int = 1, compact: bool = False) -> int:
    root = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in root.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in root.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel" if levels > 1 else "singleLevel")
    abstract.append(multi)

    for level in range(levels):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl.append(fmt)
        lvl_text = OxmlElement("w:lvlText")
        if kind == "bullet":
            lvl_text.set(qn("w:val"), "•")
        elif kind == "heading":
            lvl_text.set(qn("w:val"), ".".join(f"%{i + 1}" for i in range(level + 1)))
        else:
            lvl_text.set(qn("w:val"), f"%{level + 1}.")
        lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        if kind == "heading":
            p_style = OxmlElement("w:pStyle")
            p_style.set(qn("w:val"), f"Heading{level + 1}")
            lvl.append(p_style)
            left, hanging = 0, 0
        else:
            left = (540 if compact else 720) + level * 360
            hanging = 270 if compact else 360
        p_pr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        lvl.append(p_pr)
        abstract.append(lvl)
    root.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    root.append(num)
    return num_id


def apply_num(paragraph, num_id: int, level: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


class DocBuilder:
    def __init__(self, title: str, doc_no: str, preset: str):
        self.doc = Document()
        self.title = title
        self.doc_no = doc_no
        self.preset = preset
        self.compact = preset == "compact_reference_guide"
        self.heading_num = add_numbering(self.doc, "heading", levels=3, compact=self.compact)
        self.bullet_num = add_numbering(self.doc, "bullet", compact=self.compact)
        self.number_num = add_numbering(self.doc, "number", compact=self.compact)
        self.heading_counters = [0, 0, 0]
        self.heading_entries: list[tuple[int, str]] = []
        self.toc_paragraph = None
        self._setup()

    def _setup(self) -> None:
        doc = self.doc
        section = doc.sections[0]
        section.top_margin = Inches(0.82)
        section.bottom_margin = Inches(0.82)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.38)
        section.footer_distance = Inches(0.38)
        section.different_first_page_header_footer = True

        normal = doc.styles["Normal"]
        normal.font.name = FONT_LATIN
        normal.font.size = Pt(10.5 if self.compact else 11)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK_BODY)
        normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        normal._element.rPr.rFonts.set(qn("w:cs"), FONT_LATIN)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25 if self.compact else 1.10
        normal.font.color.rgb = rgb(INK)

        style_tokens = {
            1: (16, INK, 18 if self.compact else 16, 10 if self.compact else 8),
            2: (13, INK, 14 if self.compact else 12, 7 if self.compact else 6),
            3: (12, INK, 10 if self.compact else 8, 5 if self.compact else 4),
        }
        for level, (size, color, before, after) in style_tokens.items():
            style = doc.styles[f"Heading {level}"]
            style.font.name = FONT_LATIN
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = rgb(color)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK_HEADING)
            style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
            style._element.rPr.rFonts.set(qn("w:cs"), FONT_LATIN)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True
            if level == 1:
                style.paragraph_format.page_break_before = True

        caption = doc.styles["Caption"]
        caption.font.name = FONT_LATIN
        caption.font.size = Pt(9)
        caption.font.color.rgb = rgb(INK)
        caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK_KAI)
        caption._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        caption._element.rPr.rFonts.set(qn("w:cs"), FONT_LATIN)
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_before = Pt(4)
        caption.paragraph_format.space_after = Pt(8)

        settings = doc.settings._element
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)

        header = section.header
        hp = header.paragraphs[0]
        hp.text = f"{self.title}\t{VERSION}"
        hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        hp.paragraph_format.space_after = Pt(2)
        for run in hp.runs:
            set_east_asia(run)
            run.font.size = Pt(8.5)
            run.font.color.rgb = rgb(INK)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run(f"{TEAM}  ·  第 ")
        set_east_asia(r)
        r.font.size = Pt(8.5)
        r.font.color.rgb = rgb(INK)
        add_field(fp, " PAGE ", "1")
        r = fp.add_run(" 页")
        set_east_asia(r)
        r.font.size = Pt(8.5)
        r.font.color.rgb = rgb(INK)
        for run in fp.runs:
            set_east_asia(run)
            run.font.size = Pt(8.5)
            run.font.color.rgb = rgb(INK)

        doc.core_properties.title = self.title
        doc.core_properties.subject = PRODUCT
        doc.core_properties.author = TEAM
        doc.core_properties.last_modified_by = TEAM
        doc.core_properties.comments = "中国软件杯参赛提交文档"

    def cover(self, subtitle: str, document_type: str) -> None:
        doc = self.doc
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(54)
        r = p.add_run("中国软件杯大学生软件设计大赛 · 参赛项目技术文档")
        set_east_asia(r, FONT_CJK_KAI)
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = rgb(INK)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        picture = doc.add_picture(str(ASSET_DIR / "brand-mark.png"), width=Inches(0.9))
        picture._inline.docPr.set("descr", "云游智导产品标识")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(PRODUCT)
        set_east_asia(r, FONT_CJK_HEADING)
        r.font.size = Pt(24)
        r.font.bold = True
        r.font.color.rgb = rgb(INK)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(document_type)
        set_east_asia(r, FONT_CJK_HEADING)
        r.font.size = Pt(20)
        r.font.bold = True
        r.font.color.rgb = rgb(INK)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(62)
        r = p.add_run(subtitle)
        set_east_asia(r, FONT_CJK_KAI)
        r.font.size = Pt(11)
        r.font.color.rgb = rgb(INK)

        for label, value in (("文档编号", self.doc_no), ("文档版本", VERSION), ("编制团队", TEAM), ("发布日期", DATE)):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(f"{label}：")
            set_east_asia(r, FONT_CJK_HEADING)
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = rgb(INK)
            r = p.add_run(value)
            set_east_asia(r)
            r.font.size = Pt(10.5)
            r.font.color.rgb = rgb(INK)
        doc.add_page_break()

    def front_matter(self, purpose: str) -> None:
        p = self.doc.add_paragraph()
        r = p.add_run("文档控制")
        set_east_asia(r, FONT_CJK_HEADING)
        r.font.size = Pt(16)
        r.bold = True
        r.font.color.rgb = rgb(INK)
        p.paragraph_format.space_after = Pt(10)
        self.table(
            ["版本", "日期", "编制/审核", "变更说明"],
            [[VERSION, DATE, TEAM, "首次正式发布，依据当前交付代码、接口契约和实测结果编制。"]],
            [1100, 1700, 1900, 4660],
        )
        self.callout("文档目的", purpose, LIGHT_BLUE)
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run("目录")
        set_east_asia(r, FONT_CJK_HEADING)
        r.font.size = Pt(16)
        r.bold = True
        r.font.color.rgb = rgb(INK)
        toc = self.doc.add_paragraph()
        add_field(toc, ' TOC \\o "1-3" \\h \\z \\u ', "目录将在打开文档时自动更新")
        self.toc_paragraph = toc
        # Heading 1 carries page_break_before; avoid a duplicate blank page here.

    def heading(self, text: str, level: int = 1) -> None:
        self.heading_counters[level - 1] += 1
        for index in range(level, len(self.heading_counters)):
            self.heading_counters[index] = 0
        number = ".".join(str(value) for value in self.heading_counters[:level])
        numbered_text = f"{number} {text}"
        self.doc.add_heading(numbered_text, level=level)
        self.heading_entries.append((level, numbered_text))

    def para(self, text: str, bold_prefix: str | None = None, italic: bool = False) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.widow_control = True
        if bold_prefix and text.startswith(bold_prefix):
            r = p.add_run(bold_prefix)
            set_east_asia(r, FONT_CJK_HEADING)
            r.bold = True
            r.font.color.rgb = rgb(INK)
            r = p.add_run(text[len(bold_prefix):])
            set_east_asia(r)
        else:
            r = p.add_run(text)
            set_east_asia(r)
        r.italic = italic

    def bullets(self, items: list[str], level: int = 0) -> None:
        for item in items:
            p = self.doc.add_paragraph()
            apply_num(p, self.bullet_num, level)
            p.paragraph_format.space_after = Pt(4 if self.compact else 6)
            p.paragraph_format.line_spacing = 1.25 if self.compact else 1.10
            r = p.add_run(item)
            set_east_asia(r)

    def steps(self, items: list[str]) -> None:
        num_id = add_numbering(self.doc, "number", compact=self.compact)
        for item in items:
            p = self.doc.add_paragraph()
            apply_num(p, num_id)
            p.paragraph_format.space_after = Pt(5)
            r = p.add_run(item)
            set_east_asia(r)

    def table(self, headers: list[str], rows: list[list[str]], widths: list[int], font_size: float = 9.2) -> None:
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = "Table Grid"
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            shade(cell._tc.get_or_add_tcPr(), LIGHT_BLUE if self.compact else LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(header))
            set_east_asia(r, FONT_CJK_HEADING)
            r.bold = True
            r.font.size = Pt(font_size)
            r.font.color.rgb = rgb(INK)
        repeat_header(table.rows[0])
        for row_data in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row_data):
                p = cells[index].paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.08
                if len(str(value)) <= 12 and index != len(row_data) - 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(str(value))
                set_east_asia(r)
                r.font.size = Pt(font_size)
        set_table_geometry(table, widths)
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(1)

    def callout(self, label: str, text: str, fill: str = LIGHT_BLUE) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.right_indent = Inches(0.12)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        shade(p._p.get_or_add_pPr(), fill)
        r = p.add_run(f"{label}｜")
        set_east_asia(r, FONT_CJK_HEADING)
        r.bold = True
        r.font.color.rgb = rgb(INK)
        r = p.add_run(text)
        set_east_asia(r, FONT_CJK_KAI)
        r.font.color.rgb = rgb(INK)

    def code(self, lines: list[str]) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.right_indent = Inches(0.18)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.0
        shade(p._p.get_or_add_pPr(), "F5F7FA")
        for index, line in enumerate(lines):
            r = p.add_run(line)
            set_east_asia(r, FONT_CJK_BODY)
            r.font.size = Pt(8.8)
            r.font.color.rgb = rgb(INK)
            if index != len(lines) - 1:
                r.add_break()

    def figure(self, path: Path, caption: str, width: float = 6.25) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        run = p.add_run()
        inline = run.add_picture(str(path), width=Inches(width))
        inline._inline.docPr.set("descr", caption)
        cp = self.doc.add_paragraph(caption, style="Caption")
        cp.paragraph_format.keep_with_next = False

    def save(self, path: Path) -> None:
        if self.toc_paragraph is not None and self.toc_paragraph.runs:
            field_run = self.toc_paragraph.runs[0]
            field_xml = field_run._r
            for text_node in list(field_xml.findall(qn("w:t"))):
                field_xml.remove(text_node)
            field_end = None
            for child in field_xml:
                if child.tag == qn("w:fldChar") and child.get(qn("w:fldCharType")) == "end":
                    field_end = child
                    break
            if field_end is not None:
                for index, (level, entry) in enumerate(self.heading_entries):
                    text_node = OxmlElement("w:t")
                    text_node.set(qn("xml:space"), "preserve")
                    text_node.text = ("　" * (level - 1)) + entry
                    field_end.addprevious(text_node)
                    if index < len(self.heading_entries) - 1:
                        field_end.addprevious(OxmlElement("w:br"))
                set_east_asia(field_run, FONT_CJK_BODY)
                field_run.font.size = Pt(9.2)
                field_run.font.color.rgb = rgb(INK)
        self.doc.save(path)


def draw_wrapped(draw: ImageDraw.ImageDraw, box, text: str, fnt, fill=INK, align="center", max_chars=16):
    def line_size(value: str) -> tuple[float, int]:
        width = sum(
            draw.textlength(char, font=fnt.latin if ord(char) < 128 else fnt.cjk)
            for char in value
        )
        latin_bbox = draw.textbbox((0, 0), "Ag", font=fnt.latin)
        cjk_bbox = draw.textbbox((0, 0), "国", font=fnt.cjk)
        height = max(latin_bbox[3] - latin_bbox[1], cjk_bbox[3] - cjk_bbox[1])
        return width, height

    def draw_mixed(value: str, x: float, y: float, color: str) -> None:
        cursor = x
        for char in value:
            selected = fnt.latin if ord(char) < 128 else fnt.cjk
            draw.text((cursor, y), char, font=selected, fill=color)
            cursor += draw.textlength(char, font=selected)

    x1, y1, x2, y2 = box
    lines = []
    for part in text.split("\n"):
        lines.extend(wrap(part, max_chars) or [""])
    gap = int(fnt.size * 0.35)
    sizes = [line_size(line) for line in lines]
    heights = [height for _, height in sizes]
    total = sum(heights) + gap * (len(lines) - 1)
    y = y1 + (y2 - y1 - total) / 2
    # Submission typography uses black text throughout; fills/outlines retain color.
    color = "#000000"
    for line, (width, height) in zip(lines, sizes):
        x = x1 + (x2 - x1 - width) / 2 if align == "center" else x1 + 22
        draw_mixed(line, x, y, color)
        y += height + gap


def box(draw, xy, title, subtitle="", fill=WHITE, outline=LINE, title_color=NAVY, radius=20):
    draw.rounded_rectangle(xy, radius=radius, fill=f"#{fill}", outline=f"#{outline}", width=3)
    x1, y1, x2, y2 = xy
    if subtitle:
        draw_wrapped(draw, (x1 + 12, y1 + 12, x2 - 12, y1 + (y2-y1)*0.53), title, font(31, True), title_color, max_chars=12)
        draw_wrapped(draw, (x1 + 18, y1 + (y2-y1)*0.50, x2 - 18, y2 - 10), subtitle, font(22), MUTED, max_chars=20)
    else:
        draw_wrapped(draw, (x1 + 14, y1 + 10, x2 - 14, y2 - 10), title, font(28, True), title_color, max_chars=16)


def arrow(draw, start, end, color=ORANGE):
    draw.line([start, end], fill=f"#{color}", width=6)
    ex, ey = end
    sx, sy = start
    if abs(ex-sx) > abs(ey-sy):
        direction = 1 if ex > sx else -1
        pts = [(ex, ey), (ex-18*direction, ey-12), (ex-18*direction, ey+12)]
    else:
        direction = 1 if ey > sy else -1
        pts = [(ex, ey), (ex-12, ey-18*direction), (ex+12, ey-18*direction)]
    draw.polygon(pts, fill=f"#{color}")


def make_assets() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    mark = Image.new("RGB", (360, 360), f"#{LIGHT_ORANGE}")
    d = ImageDraw.Draw(mark)
    d.rounded_rectangle((45, 45, 315, 315), radius=62, fill=f"#{ORANGE}")
    d.polygon([(92, 240), (166, 130), (215, 192), (254, 142), (292, 240)], fill="white")
    d.ellipse((160, 92, 198, 130), fill=f"#{WHITE}")
    mark.save(ASSET_DIR / "brand-mark.png")

    img = Image.new("RGB", (1600, 820), "#FAFBFC")
    d = ImageDraw.Draw(img)
    draw_wrapped(d, (40, 20, 1560, 92), "本地单机交付部署拓扑", font(38, True), NAVY, max_chars=30)
    boxes = [
        ((70, 210, 380, 500), "浏览器客户端", "游客端 · 团长端 · 管理端\nChrome / Edge"),
        ((505, 170, 905, 540), "FastAPI 主后端", "同源静态页面 + /api\n单 Worker · 端口 8000"),
        ((1030, 120, 1500, 330), "本地持久化", "SQLite：data/app.db\n上传文件：uploads/\n景区资料：data/*.json"),
        ((1030, 430, 1500, 680), "可选外部能力", "大模型问答 · 识图/ASR\nEdge TTS · 高德地图\n未配置时明确降级"),
    ]
    for xy, t, s in boxes:
        box(d, xy, t, s, fill=WHITE, outline=LINE)
    arrow(d, (380, 355), (505, 355))
    arrow(d, (905, 300), (1030, 230), TEAL)
    arrow(d, (905, 405), (1030, 545), ORANGE)
    draw_wrapped(d, (385, 275, 505, 330), "HTTP / WS", font(20, True), MUTED, max_chars=12)
    img.save(ASSET_DIR / "deployment.png")

    img = Image.new("RGB", (1600, 980), "#FAFBFC")
    d = ImageDraw.Draw(img)
    draw_wrapped(d, (40, 20, 1560, 90), "云游智导总体逻辑架构", font(38, True), NAVY, max_chars=30)
    layers = [
        (120, 155, 1480, 300, "表现层", "身份入口｜游客随行导览｜团长工作台｜运营管理后台", LIGHT_ORANGE, ORANGE),
        (120, 345, 1480, 490, "接口与协同层", "REST API｜WebSocket 房间事件｜统一错误模型｜Bearer 认证", LIGHT_BLUE, BLUE),
        (120, 535, 1480, 700, "业务服务与智能编排层", "问答与隐私决策｜RAG 检索｜路线推荐｜识图/ASR/TTS｜运营统计", LIGHT_GREEN, TEAL),
        (120, 745, 1480, 895, "数据与基础设施层", "SQLite + FTS5｜JSON 景区资料｜本地上传｜日志/限流/迁移/备份", LIGHT_GRAY, NAVY),
    ]
    for x1, y1, x2, y2, t, s, fill, color in layers:
        d.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=f"#{fill}", outline=f"#{color}", width=3)
        draw_wrapped(d, (x1+20, y1+18, x1+300, y2-18), t, font(30, True), color, max_chars=10)
        draw_wrapped(d, (x1+330, y1+16, x2-24, y2-16), s, font(25), INK, align="left", max_chars=35)
    for y in (300, 490, 700):
        arrow(d, (800, y+4), (800, y+39), MUTED)
    img.save(ASSET_DIR / "architecture.png")

    img = Image.new("RGB", (1600, 900), "#FAFBFC")
    d = ImageDraw.Draw(img)
    draw_wrapped(d, (40, 20, 1560, 90), "可信问答与打断续讲流程", font(38, True), NAVY, max_chars=30)
    steps = [
        ("用户输入", "文字 / 语音\n当前景点与偏好"),
        ("隐私与风险决策", "公共 / 私人\n是否通知团长"),
        ("知识检索", "运营事件优先\nFTS5 + 景点回退"),
        ("模型回答", "依据资料生成\n无依据则说明边界"),
        ("语音与续讲", "暂停位置保存\n回答后自然衔接"),
    ]
    xs = [45, 355, 665, 975, 1285]
    for index, ((t, s), x) in enumerate(zip(steps, xs)):
        box(d, (x, 240, x+270, 580), t, s, fill=WHITE, outline=BLUE if index < 3 else ORANGE)
        if index < len(steps)-1:
            arrow(d, (x+270, 410), (xs[index+1], 410), TEAL if index < 2 else ORANGE)
    draw_wrapped(d, (160, 660, 1440, 790), "设计约束：来源可追溯；未知不编造；私人需求不进入公共频道；外部能力失败时保留文字结果并返回明确错误。", font(27, True), DARK_BLUE, max_chars=45)
    img.save(ASSET_DIR / "qa-flow.png")

    img = Image.new("RGB", (1600, 940), "#FAFBFC")
    d = ImageDraw.Draw(img)
    draw_wrapped(d, (40, 20, 1560, 90), "核心数据实体与关系", font(38, True), NAVY, max_chars=30)
    entities = [
        ((70, 150, 410, 345), "用户与认证", "users\nsessions\nws_tickets\nuser_profiles"),
        ((630, 140, 970, 355), "同行房间", "rooms\nroom_members\nroom_messages\ndirect_messages"),
        ((1190, 150, 1530, 345), "运营数据", "feedback\noperation_events\nscenic_operation_events"),
        ((350, 570, 700, 785), "知识库", "kb_documents\nkb_chunks\nkb_chunks_fts"),
        ((900, 570, 1250, 785), "内容配置", "avatar_settings\ndata/*.json\nuploads/*"),
    ]
    for xy, t, s in entities:
        box(d, xy, t, s, fill=WHITE, outline=LINE)
    arrow(d, (410, 245), (630, 245), TEAL)
    arrow(d, (970, 245), (1190, 245), TEAL)
    arrow(d, (800, 355), (595, 570), ORANGE)
    arrow(d, (820, 355), (1035, 570), ORANGE)
    draw_wrapped(d, (505, 405, 1095, 520), "房间是三端协同的业务主线；知识、内容配置和运营事件共同影响问答、讲解与路线。", font(25, True), DARK_BLUE, max_chars=30)
    img.save(ASSET_DIR / "data-model.png")

    img = Image.new("RGB", (1600, 660), "#FAFBFC")
    d = ImageDraw.Draw(img)
    draw_wrapped(d, (40, 20, 1560, 90), "需求—设计—测试闭环", font(38, True), NAVY, max_chars=30)
    labels = [("业务需求", LIGHT_ORANGE, ORANGE), ("功能/非功能需求", LIGHT_BLUE, BLUE), ("方案与模块设计", LIGHT_GREEN, TEAL), ("测试用例", LIGHT_GRAY, NAVY), ("验收结论", LIGHT_ORANGE, ORANGE)]
    x = 35
    for index, (label, fill, color) in enumerate(labels):
        box(d, (x, 220, x+265, 455), label, fill=fill, outline=color)
        if index < len(labels)-1:
            arrow(d, (x+265, 337), (x+315, 337), TEAL)
        x += 315
    draw_wrapped(d, (140, 500, 1460, 600), "每一项核心需求均对应设计实现与测试证据；发现缺口后回到需求或知识治理环节修正。", font(25, True), DARK_BLUE, max_chars=45)
    img.save(ASSET_DIR / "traceability.png")


def build_manual() -> Path:
    b = DocBuilder("产品部署和使用手册", "YYZD-DOC-DEP-001", "compact_reference_guide")
    b.cover("本地交付、配置、角色操作、维护与故障处理指南", "产品部署和使用手册")
    b.front_matter("指导评审人员、部署人员和最终用户在本地环境完成安装、配置、启动、角色操作、数据维护、备份恢复与故障排查。")

    b.heading("文档说明")
    b.para("本手册适用于云游智导 V2.0.0 本地交付版。系统以前后端同源方式运行：浏览器访问同一 FastAPI 服务，后端同时提供静态页面、REST API、WebSocket、上传文件和健康检查。")
    b.table(["对象", "适用内容"], [["评审/演示人员", "一键启动、访问入口、三端功能操作与验收检查。"], ["部署人员", "环境准备、配置项、手动部署、端口与运行状态检查。"], ["管理员", "知识库、数字人配置、运营事件、数据备份与安全维护。"], ["普通用户", "游客问答、识图、路线、讲解、同行小队与反馈。"]], [1900, 7460])
    b.callout("版本边界", "当前版本面向本地单机交付，SQLite、WebSocket 连接和限流状态均按单进程设计，必须以单 Worker 启动。", LIGHT_ORANGE)

    b.heading("产品概览")
    b.para("云游智导围绕景区游览全过程，连接游客、团长和景区管理员三类角色。游客获得可信问答、图片识景、个性化路线和数字人讲解；团长统一管理房间、成员和讲解节奏；管理端维护知识、数字人形象和运营数据。")
    b.table(["角色", "主要入口", "核心能力"], [["游客", "/pages/user-portal/index.html", "景区问答、语音提问、图片识景、路线推荐、自动讲解、群聊/私信、文化护照。"], ["团长", "/pages/guide-panel/index.html", "创建同行小队、成员管理、景点切换、开始/暂停/继续讲解、广播与协助处理。"], ["景区管理员", "/pages/dashboard/index.html", "运营指标、热门问题、评分反馈、知识库、数字人设置和运营事件。"]], [1300, 2900, 5160])
    b.figure(CAPTURE_DIR / "00-latest-landing.png", "图 1 统一身份入口")

    b.heading("部署架构与环境要求")
    b.figure(ASSET_DIR / "deployment.png", "图 2 本地单机交付部署拓扑")
    b.heading("硬件与软件环境", 2)
    b.table(["类别", "最低要求", "推荐配置"], [["处理器", "双核 x64", "四核及以上 x64"], ["内存", "4 GB", "8 GB 及以上"], ["磁盘", "2 GB 可用空间", "5 GB 及以上，预留上传与备份空间"], ["操作系统", "Windows 10/11；Linux/macOS 可手动部署", "Windows 11 x64"], ["Python", "3.10 及以上", "3.11/3.12"], ["浏览器", "支持 ES6、WebSocket 的现代浏览器", "最新版 Edge 或 Chrome"], ["网络", "本地页面和本地数据可离线运行", "启用外部问答、识图、ASR、地图与 Edge TTS 时需联网"]], [1500, 3500, 4360])
    b.heading("交付目录", 2)
    b.table(["目录/文件", "说明"], [["app/", "FastAPI API、业务服务、Provider、配置、数据库与中间件。"], ["frontend-v4/", "游客端、团长端、管理端页面以及公共 CSS/JavaScript/图片。"], ["data/", "SQLite 数据库、景区知识、路线、路径图和识图景点数据。"], ["uploads/", "TTS、音频、聊天媒体、知识库和数字人上传文件。"], ["tools/preflight.py", "启动前数据、配置和能力检查。"], ["start.bat", "Windows 一键创建虚拟环境、安装依赖、自检并启动。"], ["tests/", "接口、算法、安全和真实服务条件测试。"]], [2200, 7160])

    b.heading("部署前准备")
    b.steps(["将完整项目解压到具有读写权限的本地目录，建议路径中避免过长层级。", "安装 Python 3.10 或更高版本，并确认命令行执行 python --version 可返回版本号。", "如需正式账号管理，复制 .env.example 为 .env，并将 ADMIN_PASSWORD 修改为强密码。", "按需配置外部问答、图片识别、语音识别和地图密钥；未配置的能力会明确提示不可用，不会生成伪造结果。", "确认计划使用的端口未被占用，默认端口为 8000。"])
    b.heading("关键配置项", 2)
    b.table(["配置项", "默认/示例", "用途与注意事项"], [["DATABASE_PATH", "data/app.db", "SQLite 数据库路径；需要持续备份。"], ["ADMIN_USER_NAME", "admin", "首次启动时创建的管理员名称。"], ["ADMIN_PASSWORD", "必须修改", "不要使用示例值，不要提交到版本库。"], ["SESSION_TTL_SECONDS", "86400", "账号会话有效期，单位秒。"], ["GUEST_TTL_SECONDS", "43200", "访客会话有效期，单位秒。"], ["WS_TICKET_TTL_SECONDS", "60", "WebSocket 一次性票据有效期。"], ["RATE_LIMIT_ENABLED", "true", "启用认证、消息、问答、识图和上传限流。"], ["DEEPSEEK_API_KEY", "空", "配置后启用正式大模型问答。"], ["VISION_API_KEY/BASE_URL/MODEL", "空", "完整配置后启用图片识景；密钥也可供 ASR 使用。"], ["DASHSCOPE_API_KEY", "空", "可单独提供 Qwen ASR 凭据。"], ["MAP_API_KEY", "空", "配置后启用高德地图路线能力。"], ["ENABLE_ASR/TTS/VISION/RAG", "true", "对应功能开关；关闭时接口返回明确说明。"]], [2700, 1900, 4760], 8.6)
    b.callout("保密要求", ".env 中的密钥、管理员密码和公网地址不得出现在截图、日志、提交仓库或公开文档中。", LIGHT_ORANGE)

    b.heading("Windows 一键部署")
    b.steps(["在项目根目录双击 start.bat，或在命令提示符中执行 start.bat。", "首次运行时脚本创建 .venv 并安装 requirements.txt 中的依赖。", "脚本执行 tools/preflight.py，对 Python 版本、五类数据文件、知识来源和外部能力配置进行检查。", "看到 [PASS] Preflight completed 后，服务以单 Worker 启动。", "浏览器访问 http://127.0.0.1:8000/，系统将跳转至身份入口。"])
    b.code(["cd <项目目录>", "start.bat", "REM 如 8000 端口被占用：", "start.bat 8010"])
    b.heading("手动部署", 2)
    b.code(["python -m venv .venv", ".venv\\Scripts\\python -m pip install -r requirements.txt", ".venv\\Scripts\\python tools\\preflight.py", ".venv\\Scripts\\python -m uvicorn app.main:app --host 127.0.0.1 --port 8000 --workers 1"])
    b.callout("禁止配置", "不要把 --workers 设置为大于 1；开发时可以使用 --reload，但不要同时使用 --reload 与 --workers。", LIGHT_ORANGE)

    b.heading("启动验证与停止")
    b.table(["检查项", "地址/操作", "期望结果"], [["产品入口", "http://127.0.0.1:8000/", "307 跳转至身份选择页并正常显示。"], ["存活检查", "GET /health/live", "返回 status=live。"], ["就绪检查", "GET /health/ready", "返回 status=ready、database=ok、workerMode=single。"], ["接口文档", "http://127.0.0.1:8000/docs", "显示 Swagger 接口列表。"], ["OpenAPI", "http://127.0.0.1:8000/openapi.json", "返回 V2.0.0 接口描述。"]], [1700, 3700, 3960])
    b.para("停止服务时回到启动窗口按 Ctrl+C，等待进程正常退出。不得在数据库迁移、知识库重建或文件上传过程中直接结束进程。")

    b.heading("游客端使用说明")
    b.heading("进入导览与可信问答", 2)
    b.steps(["在身份入口选择“我是游客”，按页面提示创建访客会话或使用账号登录。", "进入“问导游”，确认当前景区/景点信息。", "输入文化、建筑、设施或路线问题并发送；语音按钮可在浏览器授权麦克风后录音。", "查看回答正文、资料来源和提示信息；资料不足时系统会明确说明边界。"])
    b.figure(CAPTURE_DIR / "01-trusted-qa.png", "图 3 游客端可信知识问答与回答依据")
    b.heading("图片识景", 2)
    b.steps(["打开拍照识景功能，选择 JPEG、PNG 或 WebP 图片。", "确认预览图无误后提交识别。", "查看识别名称、置信度、视觉特征、文化说明、相关景点和资料来源。", "识别服务未配置或识别失败时，根据提示更换角度或联系现场人员确认。"])
    b.figure(CAPTURE_DIR / "03-vision-building.png", "图 4 建筑识别、特征说明与来源展示")
    b.heading("路线推荐与动态调整", 2)
    b.steps(["选择游览时长、兴趣、体力、同行成员和少走路等偏好。", "点击生成路线，查看站点、预计时间、距离、休息点和匹配原因。", "出现客流、封路、天气或设施关闭事件时，查看系统给出的路线调整和时间变化。"])
    b.figure(CAPTURE_DIR / "04-route.png", "图 5 长者友好路线规划结果")
    b.heading("数字人讲解与打断续讲", 2)
    b.steps(["到达景点或由团长切换当前景点后，进入“看讲解”。", "查看数字人状态、实时字幕、知识标签、音色和播放进度。", "讲解中提出问题时，系统暂停并保存位置；回答结束后生成衔接语继续讲解。", "语音不可用时仍保留文字回答和字幕，不影响信息获取。"])
    b.figure(CAPTURE_DIR / "06-digital-guide.png", "图 6 数字人自动讲解与实时字幕")
    b.heading("同行小队与私人协助", 2)
    b.bullets(["输入团长分享的同行码加入小队，可使用群聊、私信、图片和语音消息。", "身体不适、如厕、迷路和离队等私人需求不会完整进入公共频道。", "需要人工协助时，团长只收到风险级别、位置和必要处置提醒。", "退出小队前确认重要信息已保存；房主需先转移团长身份或结束房间。"])

    b.heading("团长端使用说明")
    b.heading("创建并管理同行小队", 2)
    b.steps(["在身份入口选择“我是团长”，完成访客团长会话或账号登录。", "选择景区和路线，输入房间名称后创建小队。", "复制同行码分享给游客，等待成员加入。", "在成员列表查看在线状态、当前路线、景点进度和协助请求；必要时可移除成员或转移团长。"])
    b.figure(CAPTURE_DIR / "08-leader-room.png", "图 7 团长端小队、路线与成员状态")
    b.heading("统一讲解与应急操作", 2)
    b.bullets(["开始讲解：生成当前景点讲解并同步至房间成员。", "暂停/继续：更新房间状态并同步播放节奏。", "景点切换：房主选择新的当前景点，成员端接收房间事件。", "集合提醒：向公共频道发布广播；私人协助通过通知单独处理。", "结束导览：将房间状态改为 ended，结束后不可恢复。"])
    b.figure(CAPTURE_DIR / "09-leader-control.png", "图 8 团长端讲解控制与协助处理")

    b.heading("管理端使用说明")
    b.heading("登录与运营看板", 2)
    b.para("管理员账号由 .env 中的 ADMIN_USER_NAME 和 ADMIN_PASSWORD 在首次启动时创建。登录后可查看服务量、问答趋势、热门问题、热门景点、满意度、系统指标和游客报告。看板数据来源于实际数据库聚合，不使用估算分数。")
    b.figure(CAPTURE_DIR / "12-admin-analytics.png", "图 9 景区运营指标与热点分析")
    b.heading("知识库管理", 2)
    b.steps(["进入知识库管理页，按名称、分类或状态查看资料。", "上传 TXT、Markdown、JSON 或 PDF 文件，单文件最大 20 MB。", "确认文档状态变为 indexed，并查看分块数量。", "使用测试查询验证中文检索；资料变化后执行“重新整理”。", "删除文档前确认无业务依赖，删除将同时移除文件、元数据和索引。"])
    b.figure(CAPTURE_DIR / "13-admin-kb.png", "图 10 知识库检索、上传与重新整理")
    b.heading("数字人和运营事件", 2)
    b.bullets(["数字人：设置形象图片、声音、语速、表情以及口型/待机行为；形象上传支持 PNG、JPG、WebP。", "运营事件：发布人流、封路、天气、公告或设施关闭信息，并设置影响景点、路线和有效期。", "有效运营事件在知识检索中优先于静态资料；标记 resolved 或 expired 后不再参与回答。"])

    b.heading("数据备份、恢复与更新")
    b.heading("备份", 2)
    b.steps(["正常停止后端服务。", "复制 data/app.db、data/backups/ 和 uploads/ 至独立备份目录。", "记录备份时间、版本和操作人，避免只保留同一磁盘副本。"])
    b.heading("恢复", 2)
    b.steps(["停止后端并另行保存当前 data/app.db。", "从 data/backups/ 选择目标备份，复制为 data/app.db。", "启动系统，执行 /health/ready 和关键角色功能检查。"])
    b.heading("版本更新", 2)
    b.bullets(["更新前完整备份 .env、data/ 和 uploads/。", "覆盖代码时保留现有 .env，不要把 .env.example 当作生产配置直接替换。", "启动时数据库迁移会自动执行；检测到旧库会先通过 SQLite Backup API 备份。", "更新后重新运行 preflight、编译检查和自动化测试。"])

    b.heading("常见故障与处理")
    b.table(["现象", "可能原因", "处理方法"], [["提示未找到 Python", "未安装或 PATH 未配置", "安装 Python 3.10+，重新打开终端并执行 python --version。"], ["端口 8000 被占用", "已有进程监听", "执行 start.bat 8010，或停止占用进程。"], ["Preflight 失败", "JSON 数据缺失/格式错误", "按输出文件名恢复 data 目录，修复后重新自检。"], ["智能问答返回 503", "DEEPSEEK_API_KEY 未配置或服务不可用", "检查配置和网络；系统不会返回伪造答案。"], ["识图提示未配置", "识图三个配置项不完整", "同时配置 VISION_API_KEY、VISION_BASE_URL、VISION_MODEL。"], ["TTS 有文字无声音", "Edge TTS 网络不可达或功能关闭", "检查 ENABLE_TTS 和网络；文字回答仍可使用。"], ["WebSocket 连接失败", "票据超时或已使用", "重新调用 /api/auth/ws-ticket 获取 60 秒一次性票据。"], ["知识文档索引失败", "格式、编码或 PDF 文本提取异常", "查看文档状态和 error 字段，修正文件后重新上传。"], ["数据库被锁", "多 Worker 或异常并发写入", "恢复为单 Worker，停止多余进程后重启。"]], [2100, 2700, 4560], 8.5)

    b.heading("安全与验收清单")
    b.heading("安全操作", 2)
    b.bullets(["首次部署必须修改管理员密码；密钥只保存在 .env。", "所有长期 HTTP 认证使用 Authorization: Bearer <token>，不得把 Token 放入 URL 或业务 JSON。", "不要公开 data/app.db、uploads/、日志和备份文件。", "公网部署前应在反向代理层启用 HTTPS、访问控制、日志轮换和独立备份。", "本地版本未设计为多实例集群，不应直接用于高并发生产环境。"])
    b.heading("交付验收", 2)
    b.table(["序号", "验收项", "合格标准"], [["1", "启动前检查", "五类数据文件通过、知识来源通过、无阻断项。"], ["2", "服务可用", "/health/live 与 /health/ready 返回成功。"], ["3", "三端入口", "游客、团长、管理员均可进入对应页面。"], ["4", "核心业务", "问答、识图、路线、讲解、小队、知识库与看板按配置工作。"], ["5", "权限", "管理员接口拒绝非 admin，房间操作校验成员与房主身份。"], ["6", "数据", "数据库、上传文件、备份目录可读写且已建立备份。"], ["7", "质量", "自动化测试 46 项通过；真实外部模型验证按环境另行执行。"]], [800, 2500, 6060])
    b.heading("附录：常用地址和命令")
    b.table(["项目", "内容"], [["产品入口", "http://127.0.0.1:8000/"], ["游客端", "http://127.0.0.1:8000/pages/user-portal/index.html"], ["团长端", "http://127.0.0.1:8000/pages/guide-panel/index.html"], ["管理看板", "http://127.0.0.1:8000/pages/dashboard/index.html"], ["接口文档", "http://127.0.0.1:8000/docs"], ["运行测试", "python -m pytest -q"], ["编译检查", "python -m compileall -q app src"], ["启动前检查", "python tools/preflight.py"]], [2300, 7060])
    b.para("文档结束。", italic=True)
    path = OUT / "云游智导_产品部署和使用手册_V1.0.docx"
    b.save(path)
    return path


def build_design() -> Path:
    b = DocBuilder("产品总体设计文档", "YYZD-DOC-DES-001", "standard_business_brief")
    b.cover("系统边界、总体架构、模块、数据、接口、安全与质量设计", "产品总体设计文档")
    b.front_matter("说明云游智导的总体技术方案和关键设计决策，为评审、开发、测试、部署及后续演进提供统一基线。")

    b.heading("设计概述")
    b.para("云游智导是一套面向景区游览全流程的三端协同系统。系统把游客的问答、识图、路线和数字人讲解，与团长的房间/成员/讲解控制、景区管理端的内容和运营治理连接为闭环。")
    b.heading("设计目标", 2)
    b.bullets(["可信：景点事实优先引用知识库，资料不足时明确不确定，不伪造模型结果。", "自适应：路线、语速、字幕和休息安排结合体力、兴趣、同行成员和现场事件调整。", "懂分寸：公共问答、私人需求和团长通知采用不同信息边界。", "可交付：同源单服务、SQLite、自动迁移和一键启动，降低评审环境部署成本。", "可治理：知识、运营事件、反馈、数字人配置和服务统计均可在管理端维护。"])
    b.heading("设计范围", 2)
    b.table(["范围内", "范围外/当前约束"], [["游客端、团长端、管理端完整业务闭环", "跨景区大规模 SaaS 租户隔离"], ["本地 FastAPI 单进程部署", "多实例 WebSocket 广播与分布式限流"], ["SQLite + FTS5 + JSON 景区数据", "高并发分布式数据库和向量数据库"], ["可配置外部问答、识图、ASR、TTS、地图", "无凭据时伪造模型成功结果"], ["浏览器移动端优先界面", "原生 iOS/Android 应用"]], [4680, 4680])

    b.heading("需求与角色边界")
    b.table(["角色", "核心诉求", "系统职责", "主要数据边界"], [["游客", "快速理解景点、少查资料、获得路线和帮助", "问答、识图、路线、讲解、聊天、反馈", "私人需求不进入公共频道；仅保存结构化偏好标签。"], ["团长", "统一带队节奏、识别协助请求", "房间、成员、景点、讲解、广播、协助", "只能查看必要协助信息，不能读取游客完整私人描述。"], ["管理员", "维护内容、监控服务质量", "知识库、看板、事件、数字人设置", "受 admin 角色控制；操作影响全局内容。"]], [1100, 2500, 2900, 2860], 8.5)
    b.heading("关键业务能力", 2)
    b.bullets(["可信知识问答和来源展示。", "佛像细节、建筑和景点图片识别。", "时间/兴趣/体力/同行成员驱动的路线推荐。", "现场运营事件驱动的动态改线与知识优先级。", "数字人自动讲解、音色/语速、字幕与打断续讲。", "同行房间、群聊、私信、媒体消息和实时状态同步。", "运营指标、知识管理、数字人配置、游客反馈和知识缺口闭环。"])

    b.heading("总体架构")
    b.figure(ASSET_DIR / "architecture.png", "图 1 云游智导总体逻辑架构")
    b.heading("架构风格", 2)
    b.para("系统采用浏览器前端 + FastAPI 模块化单体 + SQLite/本地文件的架构。前端静态资源与 API 由同一进程、同一端口提供，避免跨域联调复杂度；API、业务服务、Provider 和数据访问按目录分层。")
    b.table(["层次", "主要组件", "职责"], [["表现层", "landing、user-portal、guide-panel、dashboard、knowledge-base、avatar-studio", "完成三类角色交互和状态展示。"], ["接口层", "app/api/*、WebSocket", "路由、认证依赖、限流、Schema 校验、错误响应。"], ["业务层", "app/services/*、algorithm_facade", "问答、讲解、隐私分流、路线、知识、房间、消息、统计。"], ["Provider 层", "app/providers/*", "封装 LLM、视觉、音频和地图外部能力。"], ["核心层", "app/core/*、middleware/*", "配置、数据库迁移、日志、错误、限流和安全头。"], ["数据层", "SQLite、FTS5、data/*.json、uploads/*", "结构化数据、全文检索、静态景区数据与文件持久化。"]], [1500, 3600, 4260], 8.7)

    b.heading("部署与运行设计")
    b.figure(ASSET_DIR / "deployment.png", "图 2 本地交付部署拓扑")
    b.heading("运行约束", 2)
    b.bullets(["单 Worker：WebSocket 连接表和限流窗口位于进程内，SQLite 写入策略也按单实例设计。", "同源服务：/assets、/pages、/uploads 和 /api 由 127.0.0.1:8000 提供。", "启动生命周期：初始化数据库、创建管理员、同步内置知识、清理 .part 临时文件，并每 300 秒清理过期认证状态。", "数据迁移：使用 schema_migrations 记录版本；旧库或版本升级前通过 SQLite Backup API 生成备份。"])
    b.heading("技术栈", 2)
    b.table(["类别", "技术", "选型理由"], [["前端", "HTML5、CSS、JavaScript、Tailwind 生成样式", "无需构建即可运行，移动端优先，部署简单。"], ["后端", "Python 3.10+、FastAPI、Uvicorn、Pydantic v2", "类型清晰、OpenAPI 自动生成、异步外部调用。"], ["数据", "SQLite WAL、FTS5 trigram、JSON", "本地交付零运维、中文全文检索、景区数据可版本化。"], ["实时", "WebSocket", "房间消息、成员、景点和状态实时同步。"], ["音频", "Edge TTS、可配置 ASR Provider", "保留文字降级，统一 /uploads 稳定 URL。"], ["测试", "pytest、FastAPI TestClient", "覆盖接口、算法、安全、迁移和异常降级。"]], [1400, 3300, 4660])

    b.heading("功能模块设计")
    b.heading("游客随行导览模块", 2)
    b.bullets(["问导游：文字/浏览器语音输入，带当前景点和偏好上下文。", "拍照识景：验证图片格式和大小，返回识别对象、视觉特征、文化说明与来源。", "路线推荐：对路线模板进行偏好评分，返回站点、时间、距离、原因和地图信息。", "数字人讲解：统一音色、语速、字幕、音频状态，支持到点自动触发。", "打断续讲：保存原讲解位置，回答完成后生成一句衔接语并恢复。", "同行与隐私：公共消息、私信、媒体消息、私人需求和团长提醒分流。"])
    b.heading("团长协同模块", 2)
    b.bullets(["房间生命周期：创建、加入、离开、移除、转移团长、暂停和结束。", "成员与路线：集中展示成员状态、当前路线、当前景点和站点进度。", "讲解控制：房主启动讲解、切换景点、暂停/继续并同步数字人状态。", "实时协同：一次性票据建立 WebSocket，广播成员、消息、景点、团长和房间状态。"])
    b.heading("景区运营管理模块", 2)
    b.bullets(["知识库：文档上传、提取、分块、FTS5 索引、查询、重建和删除。", "运营看板：基于实际数据库聚合服务量、热门问题、评分和系统指标。", "数字人配置：服务端持久化形象、声音、语速、表情和行为开关。", "运营事件：时效化的人流、封路、天气、公告和设施关闭信息进入知识优先队列。"])

    b.heading("智能服务设计")
    b.figure(ASSET_DIR / "qa-flow.png", "图 3 可信问答与打断续讲流程")
    b.heading("知识检索与来源治理", 2)
    b.para("知识检索先合并当前有效的运营事件，再查询 SQLite FTS5；完整短语无法命中时，按当前景点回退，并以中文子串/短语重叠补充。每条结果携带 title、chunkId、source、score 和 contentPreview，回答只引用实际检索结果。")
    b.bullets(["内置景区知识具有 sourceTier 与 provenance，启动前检查阻止缺失来源的数据进入正式运行。", "管理员上传文档与内置知识分离，内置数据同步不会删除上传内容。", "未检索到直接依据时，响应携带警告并要求模型说明不确定。", "运营事件在有效期内优先，关闭或过期后自动退出检索。"])
    b.heading("外部能力与降级", 2)
    b.table(["能力", "正式 Provider", "失败/未配置行为"], [["大模型问答", "DeepSeek 兼容接口", "返回 503，不返回模板化成功答案。"], ["图片识景", "配置化 Vision Provider", "返回“未配置/暂不可用”、空识别结果。"], ["语音识别", "Qwen ASR/兼容 Provider", "保留录音或文字入口，返回明确错误。"], ["语音合成", "Edge TTS/兼容 Provider", "保留文字答案并提示语音暂不可播放。"], ["地图", "高德地图", "回退到本地路线节点与提示信息。"]], [1600, 2800, 4960])

    b.heading("数据设计")
    b.figure(ASSET_DIR / "data-model.png", "图 4 核心数据实体与关系")
    b.heading("核心表", 2)
    b.table(["实体", "用途", "关键关系/约束"], [["users / sessions", "账号、访客与会话", "用户名规范化唯一；会话只存 Token 哈希并有过期时间。"], ["rooms / room_members", "同行小队和成员", "房间关联团长、路线、当前景点和状态；成员唯一。"], ["room_messages", "公共消息和广播", "按 room_id 与时间游标查询，支持文本和媒体。"], ["direct_messages / conversation_reads", "私信与已读状态", "限定房间成员和对端用户。"], ["kb_documents / kb_chunks / FTS", "上传资料、分块和全文索引", "删除文档级联移除分块和 FTS 条目。"], ["feedback", "评分、场景、评论和标签", "按用户和房间更新，供看板聚合。"], ["operation_events", "运营时效事件", "影响景点/路线和有效期，支持 resolved/expired。"], ["avatar_settings", "数字人全局配置", "管理员写、所有端读取。"], ["user_profiles", "结构化偏好记忆", "只保存标签，不保存原始私人对话。"]], [2300, 2700, 4360], 8.3)
    b.heading("文件存储", 2)
    b.table(["路径", "内容", "处理策略"], [["data/app.db", "业务数据库", "WAL 模式、自动迁移、版本升级前备份。"], ["data/*.json", "景区知识、路线、节点、边、识图条目", "启动前 JSON 校验和来源准入。"], ["uploads/audio", "游客音频", "扩展名、MIME、文件签名和大小校验，随机文件名原子落盘。"], ["uploads/tts", "语音合成结果", "验证文件存在且非空后才返回 URL。"], ["uploads/kb", "知识文档", "限制 TXT/MD/JSON/PDF 和 20 MB，索引失败记录错误。"], ["data/backups", "SQLite 迁移备份", "按时间命名，恢复时必须停止服务。"]], [1800, 3100, 4460])

    b.heading("接口与实时通信设计")
    b.heading("接口分组", 2)
    b.table(["分组", "基础路径", "职责"], [["认证", "/api/auth", "注册、登录、访客、当前用户、登出、WS 票据。"], ["房间", "/api/rooms", "房间生命周期、成员、团长、景点、讲解状态。"], ["消息", "/api/rooms/{id}/messages|direct", "公共消息、私信、会话和媒体上传。"], ["智能服务", "/api/ai、/api/audio、/api/vision", "问答、语音上传/识别/合成和图片识别。"], ["路线与景点", "/api/recommend、/api/routes、/api/spots、/api/map", "路线评分、目录、附近景点和地图。"], ["管理", "/api/kb、/api/dashboard、/api/avatar-settings、/api/operation-events", "内容、运营、数字人和事件治理。"]], [1500, 3400, 4460])
    b.heading("统一协议", 2)
    b.bullets(["除公开接口外使用 Authorization: Bearer <token>；Token 不进入 URL、FormData 或业务 JSON。", "统一错误结构为 detail、errorCode、requestId；Provider 异常使用 502/503，限流使用 429 + Retry-After。", "所有请求由 Pydantic Schema 校验长度、范围和枚举；文件上传执行独立内容校验。", "WebSocket 连接前换取 60 秒一次性票据；票据与房间、用户绑定，使用后失效。"])
    b.heading("实时事件", 2)
    b.table(["事件", "用途"], [["room.connected", "连接确认与当前房间状态。"], ["room.message", "公共消息或广播。"], ["room.members", "成员列表变化。"], ["room.leader", "团长转移。"], ["room.spot", "当前景点切换。"], ["room.status", "active/paused/ended 状态变化。"], ["leader_notify", "私人风险的必要协助提醒。"], ["pong/error", "心跳与错误。"]], [2600, 6760])

    b.heading("安全与隐私设计")
    b.table(["控制域", "设计措施"], [["密码", "PBKDF2-HMAC-SHA256，随机 16 字节盐，310,000 次迭代；恒定时间比较。"], ["会话", "随机 URL-safe Token，数据库仅保存 SHA-256 哈希；过期与登出即时清理。"], ["权限", "tourist/guide/admin 角色和房间成员/房主双重校验。"], ["实时连接", "60 秒一次性 WebSocket 票据，绑定用户和房间，使用后标记。"], ["隐私", "公共/私人决策、私人消息独立存储、团长只接收必要提醒；偏好仅保存结构化标签。"], ["上传", "限制请求体、文件类型、MIME、签名、解码大小；使用随机名和 .part 原子写入。"], ["Web 安全", "nosniff、DENY frame、no-referrer、no-store 安全响应头。"], ["滥用防护", "认证、消息、上传、问答、识图、路线和 TTS 分桶限流。"], ["错误", "稳定错误码和 requestId；服务端不回传异常堆栈。"]], [1800, 7560], 8.7)

    b.heading("界面与交互设计")
    b.para("界面采用移动端优先的“暖色画册”风格：暖白背景、白色卡片、低对比描边和克制的暖橙色交互强调。标题使用中文衬线风格，正文与控件使用高可读无衬线字体。")
    b.table(["设计原则", "实现"], [["清晰层级", "标题、状态、主要操作和来源依据分区呈现。"], ["最少切换", "游客端把问答、讲解和同行入口集中在底部导航。"], ["状态可见", "加载、讲解、暂停、同步、错误和降级均有文字状态。"], ["无障碍适配", "长者友好模式调整语速、字幕、路线强度和休息提醒。"], ["可信表达", "答案旁显示来源，未知内容使用谨慎提示。"], ["隐私可理解", "私人服务和公共频道以明显视觉提示区分。"]], [2200, 7160])

    b.heading("可靠性、可维护性与演进")
    b.bullets(["数据库连接启用 foreign_keys、busy_timeout、synchronous=NORMAL 和 WAL。", "启动时清理未完成的 .part 文件；外部调用配置超时和有限重试。", "日志包含 requestId 和服务计时，便于关联用户错误与后端日志。", "前端只依赖稳定 /api 契约，历史 /v1 算法服务已归档，不属于 V4 契约。", "Provider 通过工厂封装，便于替换 LLM、视觉、音频和地图实现。", "后续面向高并发时，应把 SQLite、限流和 WebSocket 状态迁移至外部数据库/缓存与消息总线。"])
    b.heading("质量验证", 2)
    b.table(["验证项", "结果"], [["启动前检查", "Python 3.12.4；88 条知识、5 条路线、22 个节点、25 条边、16 个识图条目通过；知识来源检查通过。"], ["编译检查", "python -m compileall -q app src 通过。"], ["自动化测试", "46 passed，1 skipped，耗时 19.61 秒。"], ["条件跳过项", "REAL_VALIDATION_BASE_URL 未配置时跳过真实外部 Provider 验证；该项需对已部署主后端单独执行。"]], [2600, 6760])
    b.heading("附录：设计依据")
    b.bullets(["README.md：本地交付、认证、数据和能力降级说明。", "docs/API.md：主后端 API 契约。", "docs/SOURCE_DATA_QUALITY.md：景区资料来源等级与准入规则。", "app/main.py、app/core、app/api、app/services：实际运行架构和实现。", "tests/：接口、算法、安全、迁移和真实服务条件测试。", "frontend-v4/pages/*/DESIGN.md：界面设计系统与页面规范。"])
    b.para("文档结束。", italic=True)
    path = OUT / "云游智导_产品总体设计文档_V1.0.docx"
    b.save(path)
    return path


def build_standard() -> Path:
    b = DocBuilder("软件系统需求、方案与测试说明书", "YYZD-DOC-SYS-001", "standard_business_brief")
    b.cover("需求分析、方案设计、测试说明、需求追踪与验收结论", "软件系统需求、方案与测试说明书")
    b.front_matter("按照软件系统标准文档规范，对业务需求、功能与非功能需求、总体方案、测试策略、测试结果和验收结论进行完整说明，并建立需求到测试的可追踪关系。")

    b.heading("引言")
    b.heading("编写目的", 2)
    b.para("本文档是云游智导软件系统的标准说明书，用于统一记录系统为什么建设、需要解决什么问题、采用什么方案、如何验证以及当前版本是否满足交付条件。")
    b.heading("项目背景", 2)
    b.para("景区游览中存在讲解信息分散、临时路线变化难同步、团队带队工具割裂、私人需求易被公共播报、运营数据和知识维护脱节等问题。项目以灵山胜境为重点场景，设计游客—团长—管理员三端协同闭环。")
    b.heading("术语", 2)
    b.table(["术语", "说明"], [["RAG", "检索增强生成；回答前检索景区知识与有效运营事件。"], ["ASR/TTS", "语音识别/语音合成。"], ["同行房间", "由团长创建、游客加入的实时协同空间。"], ["知识分块", "从内置资料或上传文档中提取的可检索文本单元。"], ["运营事件", "具有有效期的人流、封路、天气、公告或设施关闭信息。"], ["一次性票据", "用于建立 WebSocket 的 60 秒短期凭证，使用后失效。"]], [1700, 7660])

    b.heading("需求分析")
    b.heading("业务需求", 2)
    b.table(["编号", "业务需求", "验收目标"], [["BR-01", "游客在一个入口获得景点理解、路线和现场帮助。", "完成问答、识图、路线、讲解与同行服务闭环。"], ["BR-02", "团长掌握成员和讲解节奏。", "可建队、管理成员、切换景点、控制讲解和处理提醒。"], ["BR-03", "景区管理人员可维护知识并观察服务质量。", "提供知识库、运营看板、数字人和事件管理。"], ["BR-04", "系统避免无依据回答和隐私误播。", "来源可追溯，未知说明边界，私人问题不进入公共频道。"], ["BR-05", "评审环境可以低成本部署。", "Windows 一键启动；本地单进程、SQLite、同源访问。"]], [1000, 4000, 4360])
    b.heading("功能需求", 2)
    functional = [
        ["FR-01", "身份与认证", "支持账号注册/登录、游客/团长访客会话、管理员账号和登出。", "P0"],
        ["FR-02", "可信问答", "结合当前景点检索知识，返回回答、来源和警告。", "P0"],
        ["FR-03", "语音问答", "支持音频上传、ASR、公共/私人判断和回答语音。", "P1"],
        ["FR-04", "图片识景", "支持图片上传/URL、景点或细节识别、特征和来源展示。", "P0"],
        ["FR-05", "路线推荐", "按时间、兴趣、体力、同行成员和偏好推荐路线并解释原因。", "P0"],
        ["FR-06", "动态改线", "有效运营事件影响知识和路线建议，过期后退出。", "P1"],
        ["FR-07", "数字人讲解", "生成当前景点讲解音频、字幕和状态，支持语速/音色。", "P0"],
        ["FR-08", "打断续讲", "暂停并保存讲解位置，回答后生成衔接语继续。", "P1"],
        ["FR-09", "同行房间", "创建、加入、退出、移除、转移团长和结束。", "P0"],
        ["FR-10", "实时消息", "公共消息、广播、私信、图片/语音媒体和 WebSocket 同步。", "P0"],
        ["FR-11", "私人协助", "私人问题仅用户可见，必要时向团长发送最小提醒。", "P0"],
        ["FR-12", "知识库", "管理员上传、列表、详情、删除、重建和测试查询。", "P0"],
        ["FR-13", "运营看板", "展示基于数据库聚合的服务量、热点、评分和系统指标。", "P1"],
        ["FR-14", "数字人配置", "管理员设置形象、声音、语速、表情和动作行为。", "P1"],
        ["FR-15", "用户反馈", "房间成员提交 1—5 分评分、场景、标签和评论。", "P1"],
    ]
    b.table(["编号", "名称", "需求描述", "优先级"], functional, [1000, 1500, 5860, 1000], 8.2)
    b.heading("非功能需求", 2)
    b.table(["编号", "类别", "要求"], [["NFR-01", "安全", "密码不可明文；长期 Token 不落明文；角色和房间权限必须校验。"], ["NFR-02", "隐私", "私人问题不进入公共消息；只保存结构化偏好标签。"], ["NFR-03", "可靠性", "迁移前备份；未完成上传自动清理；外部失败返回明确错误。"], ["NFR-04", "性能", "本地页面快速打开；接口限流防滥用；地图结果具备缓存与请求间隔。"], ["NFR-05", "兼容性", "现代 Edge/Chrome；Python 3.10+；Windows 一键部署。"], ["NFR-06", "可维护性", "API/服务/Provider/核心分层；OpenAPI 作为字段级契约。"], ["NFR-07", "可用性", "移动端优先，状态可见，长者友好模式，语音失败保留文字。"], ["NFR-08", "可追溯性", "错误包含 requestId；知识结果包含 title/chunkId/source。"]], [1100, 1500, 6760], 8.5)
    b.heading("约束与假设", 2)
    b.bullets(["当前交付版采用单 Worker；不承诺多实例一致性。", "外部问答、识图、ASR、TTS 和地图的可用性取决于配置、网络和服务商。", "本地知识、路线、页面、管理和数据能力不依赖外部模型成功。", "景点开放时间、设施位置和运营事件存在时效性，系统必须显示来源并允许现场确认。"])

    b.heading("方案设计")
    b.figure(ASSET_DIR / "architecture.png", "图 1 软件系统总体方案")
    b.heading("方案原则", 2)
    b.bullets(["单一产品入口：静态页面、API、WebSocket 和上传文件统一由 FastAPI 提供。", "模块化单体：以清晰分层换取本地部署简洁，同时保留 Provider 可替换性。", "真实失败可见：未配置或外部失败不返回伪造成功。", "数据本地优先：核心业务、知识、路线和统计落在 SQLite/JSON/上传目录。", "安全与隐私前置：认证、权限、通道决策、上传校验和限流位于业务入口。"])
    b.heading("功能方案映射", 2)
    b.table(["需求组", "解决方案", "主要实现位置"], [["认证与权限", "Bearer 会话、PBKDF2 密码、角色/房间依赖、一次性 WS 票据", "app/services/users.py、app/api/users.py、auth dependencies"], ["问答与隐私", "algorithm_facade 决策 + RAG + Provider + 事件输出", "app/services/ai.py、knowledge.py、algorithm_facade.py"], ["识图与语音", "Provider 工厂、文件验证、稳定上传 URL、文字降级", "app/api/audio.py、vision.py；app/providers"], ["路线与事件", "路线模板评分、本地路径图、时效运营事件", "app/services/recommend.py、operation_events.py、data/*.json"], ["团队协同", "房间状态机、成员权限、公共/私信持久化和 WS 广播", "app/api/rooms.py、messages.py；app/services/rooms.py"], ["管理治理", "知识文档索引、看板聚合、数字人配置、反馈", "app/api/kb.py、dashboard.py、avatar_settings.py、feedback.py"]], [1800, 3800, 3760], 8.4)
    b.heading("关键流程", 2)
    b.figure(ASSET_DIR / "qa-flow.png", "图 2 可信问答与续讲方案")
    b.figure(ASSET_DIR / "traceability.png", "图 3 需求、设计、测试与验收闭环")

    b.heading("接口、数据与安全方案")
    b.heading("接口方案", 2)
    b.para("REST API 统一使用 /api 前缀并通过 Pydantic Schema 约束字段。公开接口限于认证入口及公共景点/路线目录；其他接口使用 Bearer Token。WebSocket 使用单独一次性票据，避免长期 Token 暴露在连接 URL 中。")
    b.table(["类别", "控制措施"], [["错误", "统一 detail/errorCode/requestId；429 携带 Retry-After；外部服务错误用 502/503。"], ["数据校验", "文本长度、数值范围、枚举、文件大小、MIME 和签名均在入口校验。"], ["授权", "角色、房间成员、房主、私信对端和管理员操作逐层校验。"], ["实时", "ticket 绑定 roomId/userId，有效 60 秒且只能使用一次。"]], [1700, 7660])
    b.heading("数据方案", 2)
    b.figure(ASSET_DIR / "data-model.png", "图 4 核心数据模型")
    b.para("业务数据使用 SQLite WAL；中文知识使用 FTS5 trigram；景区基础数据使用可审计 JSON；文件上传按用途分目录。数据库迁移按版本执行，并在旧库或升级前生成备份。")
    b.heading("安全方案", 2)
    b.bullets(["PBKDF2-HMAC-SHA256 + 随机盐 + 310,000 次迭代。", "会话 Token 和 WebSocket 票据仅以 SHA-256 哈希持久化。", "限流：认证 10/分钟、AI 30/分钟、识图 20/分钟、路线 20/分钟、消息 60/分钟、上传按 10 分钟窗口控制。", "请求体、音频、图片、知识文档和聊天媒体分别限制大小与类型。", "返回 nosniff、X-Frame-Options: DENY、no-referrer、no-store。"])

    b.heading("测试说明")
    b.heading("测试目标与范围", 2)
    b.para("测试验证接口契约、认证与权限、房间生命周期、实时消息、上传安全、外部能力失败可见、知识检索、路线推荐、数字人配置、反馈统计、数据库迁移、隐私分流和安全告警。测试不以虚构外部 Provider 成功结果替代真实服务验证。")
    b.heading("测试环境", 2)
    b.table(["项目", "实际环境"], [["操作系统", "Windows x64"], ["Python", "3.12.4"], ["测试框架", "pytest + FastAPI TestClient"], ["数据库", "测试隔离 SQLite；迁移测试使用临时目录"], ["执行命令", "python -m pytest -q"], ["执行结果", "46 passed，1 skipped，19.61 秒"], ["跳过原因", "未配置 REAL_VALIDATION_BASE_URL；真实外部 Provider 测试需对部署环境单独执行。"]], [2300, 7060])
    b.heading("测试类型", 2)
    b.table(["测试类型", "覆盖内容"], [["单元/算法测试", "公共问答中断、私人需求不广播、安全升级、未知不编造、识图与 RAG、路线偏好、记忆标签、ASR 低置信度。"], ["接口集成测试", "认证、房间、成员、消息、WebSocket 票据、音频、TTS、识图、反馈、看板、知识库、数字人和 OpenAPI。"], ["安全测试", "密码/Token 存储、身份伪造、限流、文件签名、伪造 Base64、错误码和安全响应头。"], ["数据测试", "迁移幂等、旧库/版本库备份、失败回滚、内置知识同步且保留上传文档。"], ["条件真实测试", "通过已部署主后端 /api 链路连接真实 Provider，不使用内部 /v1。"]], [2100, 7260], 8.5)
    b.heading("代表性测试用例", 2)
    cases = [
        ["TC-01", "错误认证", "使用缺失/错误 Bearer Token 访问受保护接口", "返回 401/403，稳定错误码和 requestId", "通过"],
        ["TC-02", "房间生命周期", "创建、加入、转移团长、退出、结束", "权限与状态机符合规则，ended 不可恢复", "通过"],
        ["TC-03", "一次性票据", "同一 WS ticket 连接两次", "第一次成功，第二次拒绝", "通过"],
        ["TC-04", "伪造音频", "扩展名与文件签名不匹配", "415/422，文件不落盘", "通过"],
        ["TC-05", "模型未配置", "调用正式问答但无 Provider", "503，不返回 Mock 成功", "通过"],
        ["TC-06", "可信问答", "有知识与无直接知识两类问题", "有来源；无依据显示警告且不编造", "通过"],
        ["TC-07", "私人需求", "公共房间中提出身体不适", "不持久化为公共内容；团长仅收必要提醒", "通过"],
        ["TC-08", "知识库", "中文文档上传、查询、重建、删除", "状态和分块正确，FTS 可检索，删除清理完整", "通过"],
        ["TC-09", "数据库迁移", "旧库升级和失败迁移", "升级前备份；失败回滚", "通过"],
        ["TC-10", "看板聚合", "提交反馈和问答后读取看板", "指标来自实际数据库，不显示估算评价", "通过"],
    ]
    b.table(["编号", "场景", "操作", "期望结果", "结果"], cases, [900, 1500, 2600, 3460, 900], 7.8)
    b.heading("启动前与静态质量检查", 2)
    b.table(["检查", "结果"], [["data/scenic_chunks.json", "88 records，通过"], ["data/routes.json", "5 records，通过"], ["data/path_nodes.json", "22 records，通过"], ["data/path_edges.json", "25 records，通过"], ["data/vision_spots.json", "16 records，通过"], ["知识来源治理", "通过，无缺失 sourceTier"], ["Python 编译", "app、src compileall 通过"], ["管理员密码", "当前环境自检提示需使用非示例强密码；部署前必须处理"]], [3200, 6160])

    b.heading("需求追踪与验收")
    b.heading("需求—设计—测试追踪矩阵", 2)
    trace = [
        ["FR-01", "认证服务、角色依赖、会话表", "TC-01、认证接口测试", "满足"],
        ["FR-02", "AI 服务、知识检索、来源返回", "TC-05、TC-06", "满足"],
        ["FR-03", "音频上传、ASR/TTS Provider", "音频签名、格式、低置信度测试", "满足；真实 Provider 条件验证"],
        ["FR-04", "Vision Provider + 知识补充", "伪造图片、全部演示景点识别测试", "满足；真实 Provider 条件验证"],
        ["FR-05/06", "路线评分、运营事件", "路线偏好、记忆标签、事件测试", "满足"],
        ["FR-07/08", "讲解生成、位置保存、resumeText", "讲解音频和自然衔接测试", "满足"],
        ["FR-09/10", "房间、消息、WebSocket", "TC-02、TC-03、消息游标测试", "满足"],
        ["FR-11", "公共/私人决策与 leader_notify", "TC-07、安全告警测试", "满足"],
        ["FR-12", "知识文档、分块、FTS5", "TC-08、知识同步测试", "满足"],
        ["FR-13/15", "统计、反馈和看板接口", "TC-10", "满足"],
        ["FR-14", "avatar_settings 持久化和 admin 权限", "数字人设置接口测试", "满足"],
    ]
    b.table(["需求", "设计实现", "测试证据", "结论"], trace, [1400, 3000, 3300, 1660], 8.2)
    b.heading("验收结论", 2)
    b.callout("结论", "当前代码基线通过启动前数据检查、Python 编译检查和 46 项自动化测试，核心三端功能、安全、数据迁移和失败降级具备可交付证据。1 项真实外部 Provider 验证因未配置部署地址按条件跳过，应在比赛现场或正式部署环境使用主后端 /api 链路补充执行。", LIGHT_GREEN)
    b.heading("已知限制与风险", 2)
    b.table(["风险/限制", "影响", "控制与后续计划"], [["单 Worker 架构", "不适合横向扩容", "本地交付保持单 Worker；生产演进引入 PostgreSQL/Redis/消息总线。"], ["外部 Provider 波动", "问答、识图、ASR、TTS、地图可能不可用", "超时、明确错误、文字/本地路线降级；现场提前检查网络。"], ["景区数据时效", "开放、设施和路线可能变化", "运营事件优先、来源分级、人工审核和现场确认提示。"], ["浏览器权限", "麦克风/相机授权可能失败", "提供文字输入和文件选择路径，展示授权说明。"], ["管理员示例密码", "首次部署存在弱口令风险", "启动前强制人工修改 .env；文档验收清单明确检查。"]], [2200, 2800, 4360], 8.4)

    b.heading("附录：执行命令与参考文件")
    b.code(["python tools/preflight.py", "python -m compileall -q app src", "python -m pytest -q", "python -m uvicorn app.main:app --host 127.0.0.1 --port 8000 --workers 1"])
    b.bullets(["需求与接口依据：docs/SPEC.md、docs/API.md、README.md。", "架构和实现依据：app/main.py、app/api、app/services、app/core、frontend-v4。", "数据治理依据：docs/SOURCE_DATA_QUALITY.md、data/*.json。", "测试证据：tests/、tools/preflight.py 和本次实际执行结果。"])
    b.para("文档结束。", italic=True)
    path = OUT / "云游智导_软件系统需求方案与测试说明书_V1.0.docx"
    b.save(path)
    return path


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    make_assets()
    paths = [build_manual(), build_design(), build_standard()]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
