"""Profile the optional demonstration source pack before any product import."""

from __future__ import annotations

import json
import re
from pathlib import Path

import pandas as pd
from docx import Document


SOURCE = Path("C:/Users/JRZha/Downloads/示范景区公开资料包")
OUTPUT = Path(__file__).resolve().parents[1] / "qa" / "source-pack" / "data-quality-report.json"
TARGET_SCENIC_WORDS = ("灵山", "大佛", "九龙灌浴", "梵宫", "五印坛城")


def audit_workbook(path: Path) -> dict:
    frame = pd.read_excel(path)
    numeric_columns = [
        "age", "stay_duration", "ticket_cost", "food_cost", "shopping_cost",
        "transport_cost", "entertainment_cost", "total_cost", "group_size", "satisfaction",
    ]
    spend_columns = [
        "ticket_cost", "food_cost", "shopping_cost", "transport_cost", "entertainment_cost",
    ]
    visit_date = pd.to_datetime(frame["visit_date"], errors="coerce")
    expected_total = frame[spend_columns].sum(axis=1)
    attraction_counts = frame["attraction_name"].value_counts()
    target_mask = frame["attraction_name"].astype(str).str.contains("|".join(TARGET_SCENIC_WORDS), regex=True)
    nickname_per_id = frame.groupby("tourist_id")["user_nickname"].nunique(dropna=False)
    rows_per_user = frame.groupby("tourist_id").size()

    return {
        "rows": int(len(frame)),
        "columns": list(frame.columns),
        "null_rates": {key: round(float(value), 6) for key, value in frame.isna().mean().items()},
        "exact_duplicate_rows": int(frame.duplicated().sum()),
        "duplicate_tourist_attraction_date_rows": int(
            frame.duplicated(["tourist_id", "attraction_name", "visit_date"]).sum()
        ),
        "unique_tourist_ids": int(frame["tourist_id"].nunique()),
        "unique_nicknames": int(frame["user_nickname"].nunique()),
        "tourist_ids_with_multiple_nicknames": int((nickname_per_id > 1).sum()),
        "unique_attractions": int(frame["attraction_name"].nunique()),
        "attraction_frequency": {
            "min": int(attraction_counts.min()),
            "median": float(attraction_counts.median()),
            "max": int(attraction_counts.max()),
            "coefficient_of_variation": round(float(attraction_counts.std() / attraction_counts.mean()), 6),
        },
        "rows_per_user": {
            "min": int(rows_per_user.min()),
            "median": float(rows_per_user.median()),
            "max": int(rows_per_user.max()),
            "distinct_values": sorted(int(value) for value in rows_per_user.unique())[:30],
        },
        "visit_date": {
            "invalid": int(visit_date.isna().sum()),
            "min": visit_date.min().isoformat() if visit_date.notna().any() else None,
            "max": visit_date.max().isoformat() if visit_date.notna().any() else None,
            "unique_days": int(visit_date.dt.date.nunique()),
        },
        "numeric_ranges": {
            column: {
                "min": float(frame[column].min()),
                "median": float(frame[column].median()),
                "max": float(frame[column].max()),
            }
            for column in numeric_columns
        },
        "invalid_satisfaction_rows": int((~frame["satisfaction"].between(1, 5)).sum()),
        "invalid_age_rows": int((~frame["age"].between(0, 100)).sum()),
        "invalid_group_size_rows": int((frame["group_size"] < 1).sum()),
        "total_cost_mismatch_rows": int(((frame["total_cost"] - expected_total).abs() > 0.011).sum()),
        "target_scenic_rows": int(target_mask.sum()),
        "target_scenic_share": round(float(target_mask.mean()), 6),
        "target_attractions": sorted(frame.loc[target_mask, "attraction_name"].dropna().unique().tolist()),
        "top_attractions": attraction_counts.head(10).to_dict(),
    }


def audit_docx(path: Path) -> dict:
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_rows = []
    for table in document.tables:
        for row in table.rows:
            table_rows.append([cell.text.strip() for cell in row.cells])
    incomplete_rows = sum(
        1 for row in table_rows[1:] if row and sum(not cell for cell in row) >= max(2, len(row) // 2)
    )
    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "table_rows": len(table_rows),
        "incomplete_table_rows": incomplete_rows,
        "source_url_count": len(re.findall(r"https?://", text)),
        "source_label_count": len(re.findall(r"来源|参考资料|参考文献|数据出处", text)),
        "time_sensitive_claim_count": len(re.findall(r"每日|全天开放|开放时间|票价|元/人|演出时间", text)),
        "absolute_claim_count": len(re.findall(r"唯一|第一|最大|最高|完全一致|不受天气", text)),
        "contains_lingshan": "灵山胜境" in text,
    }


def main() -> None:
    workbook = next(SOURCE.glob("*.xlsx"))
    documents = sorted(SOURCE.glob("*.docx"))
    report = {
        "sourceDirectory": str(SOURCE),
        "intendedUse": "灵山胜境知识库、路线与运营洞察候选数据源",
        "workbook": {"file": workbook.name, **audit_workbook(workbook)},
        "documents": {document.name: audit_docx(document) for document in documents},
    }
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT.write_text(json.dumps(report, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2, default=str))


if __name__ == "__main__":
    main()
