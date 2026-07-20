from __future__ import annotations

import json
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCS = sorted(
    path
    for path in (ROOT / "deliverables").glob("**/*.docx")
    if not path.name.startswith("~$")
)

EXPECTED = {
    "DEP": ["部署", "硬件与软件环境", "启动", "游客端", "团长端", "管理端", "备份", "故障排查"],
    "DES": ["总体架构", "功能模块设计", "数据设计", "接口与实时通信设计", "安全与隐私设计", "部署与运行设计", "演进"],
    "SYS": ["需求分析", "功能需求", "非功能需求", "方案设计", "测试说明", "需求追踪", "验收"],
}

LATIN_FONT = "Consolas"
ALLOWED_CJK_FONTS = {"宋体", "黑体", "楷体"}


def all_paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def effective_fonts(document: Document, paragraph, run) -> tuple[str | None, str | None]:
    ascii_font = None
    east_asia_font = None
    sources = []
    if run._r.rPr is not None:
        sources.append(run._r.rPr)
    if paragraph.style is not None and paragraph.style._element.rPr is not None:
        sources.append(paragraph.style._element.rPr)
    normal_rpr = document.styles["Normal"]._element.rPr
    if normal_rpr is not None:
        sources.append(normal_rpr)
    for rpr in sources:
        rfonts = rpr.rFonts
        if rfonts is None:
            continue
        if ascii_font is None:
            ascii_font = rfonts.get(qn("w:ascii")) or rfonts.get(qn("w:hAnsi"))
        if east_asia_font is None:
            east_asia_font = rfonts.get(qn("w:eastAsia"))
    return ascii_font, east_asia_font


def effective_color(document: Document, paragraph, run) -> str | None:
    sources = []
    if run._r.rPr is not None:
        sources.append(run._r.rPr)
    if paragraph.style is not None and paragraph.style._element.rPr is not None:
        sources.append(paragraph.style._element.rPr)
    normal_rpr = document.styles["Normal"]._element.rPr
    if normal_rpr is not None:
        sources.append(normal_rpr)
    for rpr in sources:
        color = rpr.find(qn("w:color"))
        if color is not None:
            return color.get(qn("w:val"))
    return None


def document_code(document: Document) -> str:
    comments = document.core_properties.comments or ""
    for key in EXPECTED:
        if key in comments:
            return key
    text = "\n".join(p.text for p in document.paragraphs[:20])
    if "部署和使用手册" in text:
        return "DEP"
    if "总体设计文档" in text:
        return "DES"
    return "SYS"


def audit(path: Path) -> dict[str, object]:
    with zipfile.ZipFile(path) as archive:
        bad_member = archive.testzip()
        document_xml = archive.read("word/document.xml").decode("utf-8")
        settings_xml = archive.read("word/settings.xml").decode("utf-8")

    document = Document(path)
    full_text = "\n".join(p.text for p in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            full_text += "\n" + "\t".join(cell.text for cell in row.cells)

    code = document_code(document)
    missing = [term for term in EXPECTED[code] if term not in full_text]
    replacement_characters = full_text.count("\ufffd")
    suspicious_mojibake = sum(full_text.count(token) for token in ("锟斤拷", "鏂囨", "鎴戜滑"))
    gbk_errors = sorted({char for char in full_text if not char.encode("gbk", errors="ignore")})
    headings = [p for p in document.paragraphs if p.style.name.startswith("Heading ")]
    heading_levels = [int(p.style.name.split()[-1]) for p in headings]
    hierarchy_ok = bool(heading_levels) and heading_levels[0] == 1
    previous = heading_levels[0] if heading_levels else 1
    for level in heading_levels[1:]:
        if level > previous + 1:
            hierarchy_ok = False
        previous = level

    inline_shapes = len(document.inline_shapes)
    image_alt_ok = all(
        shape._inline.docPr.get("descr")
        for shape in document.inline_shapes
    )
    table_headers_ok = True
    for table in document.tables:
        first_tr = table.rows[0]._tr
        tr_pr = first_tr.trPr
        if tr_pr is None or tr_pr.find(qn("w:tblHeader")) is None:
            table_headers_ok = False
            break

    latin_font_issues = []
    cjk_font_issues = []
    heading_font_issues = []
    caption_font_issues = []
    text_color_issues = []
    for paragraph in all_paragraphs(document):
        for run in paragraph.runs:
            if not run.text:
                continue
            ascii_font, east_asia_font = effective_fonts(document, paragraph, run)
            color = effective_color(document, paragraph, run)
            if any(ord(char) < 128 and not char.isspace() for char in run.text) and ascii_font != LATIN_FONT:
                latin_font_issues.append((run.text[:40], ascii_font))
            if any(ord(char) >= 128 for char in run.text) and east_asia_font not in ALLOWED_CJK_FONTS:
                cjk_font_issues.append((run.text[:40], east_asia_font))
            if paragraph.style.name.startswith("Heading ") and east_asia_font != "黑体":
                heading_font_issues.append((paragraph.text[:40], east_asia_font))
            if paragraph.style.name == "Caption" and east_asia_font != "楷体":
                caption_font_issues.append((paragraph.text[:40], east_asia_font))
            if color not in (None, "000000", "00000000", "auto"):
                text_color_issues.append((run.text[:40], color))

    return {
        "file": path.name,
        "size_bytes": path.stat().st_size,
        "zip_integrity": bad_member is None,
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "headings": len(headings),
        "heading_hierarchy_ok": hierarchy_ok,
        "inline_images": inline_shapes,
        "image_alt_ok": image_alt_ok,
        "table_headers_ok": table_headers_ok,
        "toc_field_present": " TOC " in document_xml,
        "update_fields_on_open": "updateFields" in settings_xml,
        "required_terms_missing": missing,
        "replacement_characters": replacement_characters,
        "suspicious_mojibake": suspicious_mojibake,
        "gbk_text_ok": not gbk_errors,
        "gbk_invalid_characters": gbk_errors,
        "latin_consolas_ok": not latin_font_issues,
        "latin_font_issues": latin_font_issues[:5],
        "cjk_font_set_ok": not cjk_font_issues,
        "cjk_font_issues": cjk_font_issues[:5],
        "heading_heiti_ok": not heading_font_issues,
        "heading_font_issues": heading_font_issues[:5],
        "caption_kaiti_ok": not caption_font_issues,
        "caption_font_issues": caption_font_issues[:5],
        "all_text_black_ok": not text_color_issues,
        "text_color_issues": text_color_issues[:5],
        "ok": all(
            [
                bad_member is None,
                path.stat().st_size > 100_000,
                len(headings) >= 7,
                hierarchy_ok,
                inline_shapes >= 1,
                image_alt_ok,
                table_headers_ok,
                " TOC " in document_xml,
                "updateFields" in settings_xml,
                not missing,
                replacement_characters == 0,
                suspicious_mojibake == 0,
                not gbk_errors,
                not latin_font_issues,
                not cjk_font_issues,
                not heading_font_issues,
                not caption_font_issues,
                not text_color_issues,
            ]
        ),
    }


def main() -> int:
    reports = [audit(path) for path in DOCS]
    print(json.dumps(reports, ensure_ascii=False, indent=2))
    return 0 if len(reports) == 3 and all(report["ok"] for report in reports) else 1


if __name__ == "__main__":
    sys.exit(main())
